import base64
import io
import re
from typing import Dict, List, Any
import pandas as pd

def extract_document_data(filename: str, file_type: str, content_base64: str) -> Dict[str, Any]:
    try:
        raw_bytes = base64.b64decode(content_base64)
    except Exception as e:
        return {"success": False, "error": f"Invalid base64 payload: {str(e)}"}

    text_preview = ""
    tables: List[Dict[str, Any]] = []
    warnings: List[str] = []

    file_type = file_type.lower().strip('.')

    if file_type == "pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(raw_bytes))
            full_text = []
            for page in reader.pages:
                txt = page.extract_text() or ""
                full_text.append(txt)
            text_preview = "\n".join(full_text)[:1500]

            # Heuristic table extraction from text patterns (tab/comma/multiple space delimited lines)
            table_rows = []
            for line in text_preview.split("\n"):
                parts = [p.strip() for p in re.split(r'\s{2,}|,|\t', line) if p.strip()]
                if len(parts) >= 3:
                    table_rows.append(parts)

            if len(table_rows) >= 3:
                headers = table_rows[0]
                rows = []
                for row_parts in table_rows[1:]:
                    row_dict = {}
                    for i, h in enumerate(headers):
                        row_dict[h] = row_parts[i] if i < len(row_parts) else ""
                    rows.append(row_dict)
                tables.append({
                    "index": 0,
                    "columns": headers,
                    "rows": rows,
                    "confidence": 0.85,
                })
        except Exception as e:
            warnings.append(f"PDF extraction notice: {str(e)}")

    elif file_type in ["docx", "doc"]:
        try:
            import docx
            doc = docx.Document(io.BytesIO(raw_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text_preview = "\n".join(paragraphs)[:1500]

            for idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text.strip() for cell in row.cells])
                if len(table_data) >= 2:
                    headers = table_data[0]
                    rows = []
                    for row_parts in table_data[1:]:
                        row_dict = {}
                        for i, h in enumerate(headers):
                            key = h if h else f"col_{i+1}"
                            row_dict[key] = row_parts[i] if i < len(row_parts) else ""
                        rows.append(row_dict)
                    tables.append({
                        "index": idx,
                        "columns": headers,
                        "rows": rows,
                        "confidence": 0.90,
                    })
        except Exception as e:
            warnings.append(f"Word document extraction notice: {str(e)}")

    if not text_preview:
        text_preview = f"Document '{filename}' uploaded successfully."

    return {
        "success": True,
        "textPreview": text_preview,
        "tables": tables,
        "warnings": warnings,
    }
